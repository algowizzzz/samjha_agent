"""Tests for the document review agent workflow."""

import os
import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Load environment variables (including ANTHROPIC_API_KEY)
from dotenv import load_dotenv
load_dotenv()

try:
    import docx  # type: ignore
except ImportError:  # pragma: no cover
    docx = None

from external.agent.doc_review_agent import DocReviewAgent  # noqa: E402


SAMPLE_MARKDOWN = """# Overview

## Scope
This policy applies to all users and systems participating in AI-based document review.

## Requirements
The program must enforce mandatory quality checks and security reviews.

## Monitoring and Reporting
Metrics are reviewed monthly with remediation tracked through the PMO.
"""


class TestDocReviewAgent(unittest.TestCase):
    """Validate the deterministic stages of the document review agent."""

    def setUp(self):
        self.agent = DocReviewAgent()

    def test_run_pipeline_on_markdown(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            source_path = Path(tmpdir) / "sample.md"
            source_path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

            # Use current API: run_phase1 takes document_path, run_id, template_id
            state = self.agent.run_phase1(
                document_path=str(source_path),
                run_id="test-run-md",
                template_id="policy_template",
            )

            # Verify file type is set in doc_meta
            self.assertEqual(state["doc_meta"]["file_type"], "md")
            # Verify structure was populated
            self.assertIn("structure", state)
            self.assertGreater(len(state["structure"]["raw_text"]), 0)
            # Verify phase1 was executed
            self.assertIn("phase1", state)
            self.assertIn(state["phase1_status"], ["success", "running", "failed"])
            # Verify no phase3 changes applied yet
            self.assertIsNone(state.get("changes", {}).get("new_raw_text"))

    @unittest.skipIf(docx is None, "python-docx not installed")
    def test_run_pipeline_on_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "sample.docx"
            document = docx.Document()
            document.add_heading("Overview", level=1)
            document.add_heading("Scope", level=2)
            document.add_paragraph("Scope paragraph.")
            document.add_heading("Requirements", level=2)
            document.add_paragraph("Requirement text.")
            document.save(str(doc_path))

            # Use current API: run_phase1 takes document_path, run_id, template_id
            state = self.agent.run_phase1(
                document_path=str(doc_path),
                run_id="test-run-docx",
                template_id="policy_template",
            )

            # Verify file type is set in doc_meta
            self.assertEqual(state["doc_meta"]["file_type"], "docx")
            # Verify structure was populated
            self.assertIn("structure", state)
            self.assertGreater(len(state["structure"]["raw_text"]), 0)
            # Verify phase1 was executed
            self.assertIn("phase1", state)
            self.assertIn(state["phase1_status"], ["success", "running", "failed"])
            # Verify no phase3 changes applied yet
            self.assertIsNone(state.get("changes", {}).get("new_raw_text"))


if __name__ == "__main__":  # pragma: no cover
    unittest.main(verbosity=2)
