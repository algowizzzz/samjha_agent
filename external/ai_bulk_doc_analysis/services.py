"""
Backend service layer for AI Bulk Doc Analysis (isolated feature).

This module handles:
- Session/document management
- Chain management
- Run execution tracking
- PDF → Markdown conversion (placeholder; real conversion will be worker-based)

All data is stored in-memory for now (can be replaced with DB later).
"""

import logging
import os
import uuid
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available. PDF conversion will fail.")

try:
    from external.platform.llm.client import get_llm_client
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    logging.warning("LLM client not available. Step execution will fail.")

logger = logging.getLogger(__name__)


@dataclass
class Document:
    doc_id: str
    session_id: str
    original_filename: str
    file_type: str
    size_bytes: int
    status: str  # QUEUED | PROCESSING | CONVERTED | ERROR
    error_code: Optional[str] = None
    error_message: Optional[str] = None
    created_at: str = ""
    updated_at: str = ""
    converted_md_path: Optional[str] = None

    def to_dict(self) -> dict:
        d = asdict(self)
        # Remove internal path, expose preview flag instead
        d.pop('converted_md_path', None)
        d['has_converted_artifact'] = self.converted_md_path is not None
        return d


@dataclass
class Chain:
    chain_id: str
    name: str
    description: str
    updated_at: str
    chain_version_id: str
    step_count: int
    valid: bool
    steps: List[Dict]  # List of step definitions

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class StepResult:
    id: str
    run_id: str
    doc_id: str
    step_index: int
    status: str  # QUEUED | RUNNING | SUCCESS | ERROR
    input_tokens: Optional[int] = None
    output_tokens: Optional[int] = None
    model: Optional[str] = None
    latency_ms: Optional[int] = None
    error_code: Optional[str] = None
    error_message: Optional[str] = None
    output_object_key: Optional[str] = None

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class Run:
    run_id: str
    session_id: str
    chain_version_id: str
    status: str  # QUEUED | RUNNING | COMPLETE | ERROR
    created_at: str
    total_input_tokens: int = 0
    total_output_tokens: int = 0
    document_ids: List[str] = None

    def __post_init__(self):
        if self.document_ids is None:
            self.document_ids = []

    def to_dict(self) -> dict:
        return asdict(self)


class BulkDocService:
    """In-memory service for Bulk Doc Analysis (replace with DB-backed service later)."""

    def __init__(self, storage_base: Optional[Path] = None):
        self.storage_base = storage_base or Path("data/ai_bulk_doc_analysis")
        self.storage_base.mkdir(parents=True, exist_ok=True)
        self.sessions: Dict[str, List[str]] = {}  # session_id -> [doc_ids]
        self.documents: Dict[str, Document] = {}
        self.chains: Dict[str, Chain] = {}
        self.runs: Dict[str, Run] = {}
        self.step_results: Dict[str, StepResult] = {}  # key: f"{run_id}:{doc_id}:{step_index}"
        self._init_mock_chains()

    def _init_mock_chains(self):
        """Initialize with a few example chains for testing."""
        chain1 = Chain(
            chain_id="chain-example-1",
            name="Example 4-Step Analysis",
            description="Demonstrates a 4-step prompt chain with R-series dependencies",
            updated_at=datetime.utcnow().isoformat() + "Z",
            chain_version_id="cv-example-1-v1",
            step_count=4,
            valid=True,
            steps=[
                {"index": 1, "required_inputs": ["R0"], "prompt": "Summarize the document into structured sections.", "description": "Extract initial structured summary"},
                {"index": 2, "required_inputs": ["R0", "R1"], "prompt": "Expand the summary with risk implications.", "description": "Enrich analysis using prior output"},
                {"index": 3, "required_inputs": ["R0", "R1"], "prompt": "Identify gaps and missing controls.", "description": "Parallel refinement step"},
                {"index": 4, "required_inputs": ["R1", "R2", "R3"], "prompt": "Produce a final consolidated assessment.", "description": "Final synthesis"},
            ],
        )
        chain2 = Chain(
            chain_id="chain-incomplete",
            name="Incomplete Chain",
            description="Chain with missing steps (for testing invalid state)",
            updated_at=datetime.utcnow().isoformat() + "Z",
            chain_version_id="cv-incomplete-v1",
            step_count=2,
            valid=False,
            steps=[],
        )
        self.chains[chain1.chain_version_id] = chain1
        self.chains[chain2.chain_version_id] = chain2

    # -------- Session Management --------
    def ensure_session(self, user_id: str, session_id: Optional[str] = None) -> str:
        """
        Get or create a session for the user.
        If session_id provided, returns that session (if exists).
        Otherwise creates a new session.
        """
        if session_id:
            # Return existing session if provided
            if session_id in self.sessions:
                return session_id
            raise ValueError(f"Session {session_id} not found or access denied")
        
        # Create new session (for in-memory, just use user_id as session_id for backward compatibility)
        session_id = f"session_{user_id}"
        if session_id not in self.sessions:
            self.sessions[session_id] = []
        return session_id
    
    def create_session(self, user_id: str, name: Optional[str] = None) -> str:
        """Create a new session for the user."""
        import uuid
        session_id = f"session_{user_id}_{uuid.uuid4().hex[:12]}"
        self.sessions[session_id] = []
        return session_id
    
    def list_sessions(self, user_id: str) -> List[Dict]:
        """List all sessions for a user, sorted by newest first."""
        import uuid
        from datetime import datetime
        sessions = []
        for session_id, doc_ids in self.sessions.items():
            if session_id.startswith(f"session_{user_id}"):
                sessions.append({
                    "session_id": session_id,
                    "user_id": user_id,
                    "created_at": datetime.utcnow().isoformat() + "Z",
                    "updated_at": datetime.utcnow().isoformat() + "Z",
                    "document_count": len([d for d in doc_ids if d in self.documents]),
                    "name": None,
                })
        # Sort by session_id (newest last for UUID-based ones)
        sessions.sort(key=lambda x: x["session_id"], reverse=True)
        return sessions
    
    def ensure_session_old(self, user_id: str) -> str:
        """Get or create a session for the user."""
        session_id = f"session_{user_id}"
        if session_id not in self.sessions:
            self.sessions[session_id] = []
        return session_id

    # -------- Document Management --------
    def create_documents(self, session_id: str, files: List, user_id: str) -> List[Document]:
        """Create document records from uploaded files."""
        docs = []
        now = datetime.utcnow().isoformat() + "Z"

        for file in files:
            doc_id = str(uuid.uuid4())
            filename = file.filename or "unknown.pdf"
            size = len(file.read())
            file.seek(0)  # Reset for later processing

            # Detect file type from extension
            from pathlib import Path
            ext = Path(filename).suffix.lower()
            file_type_map = {
                '.pdf': 'PDF',
                '.docx': 'DOCX',
                '.txt': 'TXT',
                '.md': 'MD',
                '.csv': 'CSV',
            }
            file_type = file_type_map.get(ext, 'PDF')  # Default to PDF for unknown extensions

            doc = Document(
                doc_id=doc_id,
                session_id=session_id,
                original_filename=filename,
                file_type=file_type,
                size_bytes=size,
                status="QUEUED",
                created_at=now,
                updated_at=now,
            )
            self.documents[doc_id] = doc
            self.sessions[session_id].append(doc_id)
            docs.append(doc)

            # Store file temporarily (later: move to object storage)
            doc_dir = self.storage_base / "docs" / doc_id
            doc_dir.mkdir(parents=True, exist_ok=True)
            file_path = doc_dir / filename
            # Flask file objects support .save() method
            if hasattr(file, 'save'):
                file.save(str(file_path))
            else:
                # Fallback: write bytes directly
                file.seek(0)
                with open(file_path, "wb") as f:
                    f.write(file.read())

            # Trigger conversion (for now, synchronous; later: async worker)
            self._trigger_conversion(doc_id, file_path)

        return docs

    def _trigger_conversion(self, doc_id: str, file_path: Path):
        """Trigger file conversion - supports all file types (PDF, DOCX, TXT, MD, CSV)."""
        doc = self.documents.get(doc_id)
        if not doc:
            return

        doc.status = "PROCESSING"
        doc.updated_at = datetime.utcnow().isoformat() + "Z"

        try:
            # Route based on file type - use IngestionService logic for multi-file-type support
            file_type = doc.file_type
            
            # CSV files are handled specially - no R0.md, tasks created during run
            if file_type == 'CSV':
                doc.status = "CONVERTED"
                doc.updated_at = datetime.utcnow().isoformat() + "Z"
                logger.info(f"CSV file {doc_id} marked as CONVERTED (tasks will be created during run)")
                return
            
            # Try to use IngestionService if available (handles PDF, DOCX, TXT, MD correctly)
            try:
                from .ingestion_service import IngestionService
                
                ingestion_service = IngestionService()
                
                # Create a simple ingestion profile using a dataclass-like object
                # (avoid SQLAlchemy model dependency)
                class SimpleProfile:
                    def __init__(self):
                        self.accepted_input_types = ["PDF", "DOCX", "TXT", "MD", "CSV"]
                        self.mode = "programmatic"
                        self.vision_prompt = None
                
                profile = SimpleProfile()
                
                # Ingest file based on actual file type
                r0_content, metadata = ingestion_service._ingest_programmatic(file_path, file_type)
                
                # Store converted markdown
                md_path = file_path.with_suffix(".md")
                md_path.write_text(r0_content, encoding='utf-8')
                
                doc.converted_md_path = str(md_path)
                doc.status = "CONVERTED"
                doc.updated_at = datetime.utcnow().isoformat() + "Z"
                logger.info(f"Document {doc_id} converted successfully")
                
            except (ImportError, AttributeError) as e:
                # Fallback: if IngestionService not available, handle file types directly
                logger.warning(f"IngestionService not fully available ({e}), using direct file type handling")
                
                if file_type == 'PDF':
                    if not PDFPLUMBER_AVAILABLE:
                        raise ImportError("pdfplumber is not installed. Install with: pip install pdfplumber")
                    md_content = self._convert_pdf_to_markdown(file_path)
                elif file_type in ['TXT', 'MD']:
                    md_content = file_path.read_text(encoding='utf-8')
                elif file_type == 'DOCX':
                    try:
                        from docx import Document as DocxDocument
                    except ImportError:
                        raise ImportError("python-docx not installed. Install with: pip install python-docx")
                    
                    md_parts = [f"# Document: {file_path.name}\n\n"]
                    doc_docx = DocxDocument(file_path)
                    for para in doc_docx.paragraphs:
                        if para.text.strip():
                            if para.style.name.startswith('Heading'):
                                level = para.style.name.replace('Heading ', '')
                                md_parts.append(f"{'#' * int(level)} {para.text}\n\n")
                            else:
                                md_parts.append(f"{para.text}\n\n")
                    md_content = ''.join(md_parts)
                else:
                    raise RuntimeError(f"Unsupported file type: {file_type}")
                
                # Store converted markdown
                md_path = file_path.with_suffix(".md")
                md_path.write_text(md_content, encoding='utf-8')
                
                doc.converted_md_path = str(md_path)
                doc.status = "CONVERTED"
                doc.updated_at = datetime.utcnow().isoformat() + "Z"
                logger.info(f"Document {doc_id} converted successfully")
                
        except Exception as e:
            doc.status = "ERROR"
            doc.error_code = "CONVERSION_FAILED"
            doc.error_message = str(e)
            doc.updated_at = datetime.utcnow().isoformat() + "Z"
            logger.error(f"Conversion failed for {doc_id}: {e}", exc_info=True)

    def _convert_pdf_to_markdown(self, pdf_path: Path) -> str:
        """Convert PDF file to Markdown using pdfplumber (legacy method, used as fallback only)."""
        markdown_parts = []
        markdown_parts.append(f"# Document: {pdf_path.name}\n\n")
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    markdown_parts.append(f"## Page {page_num}\n\n")
                    
                    # Extract text
                    text = page.extract_text()
                    if text:
                        # Clean up text and preserve formatting
                        text = text.strip()
                        # Convert to markdown-friendly format (double newlines for paragraphs)
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        markdown_parts.append('\n\n'.join(paragraphs))
                        markdown_parts.append('\n\n')
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, start=1):
                            markdown_parts.append(f"### Table {table_num} (Page {page_num})\n\n")
                            if table and len(table) > 0:
                                # Convert table to markdown
                                # Header row
                                if table[0]:
                                    header = "| " + " | ".join(str(cell or "") for cell in table[0]) + " |"
                                    separator = "| " + " | ".join("---" for _ in table[0]) + " |"
                                    markdown_parts.append(header)
                                    markdown_parts.append(separator)
                                
                                # Data rows
                                for row in table[1:]:
                                    if row:
                                        row_md = "| " + " | ".join(str(cell or "") for cell in row) + " |"
                                        markdown_parts.append(row_md)
                                markdown_parts.append('\n\n')
        
        except Exception as e:
            logger.error(f"Error converting PDF {pdf_path}: {e}", exc_info=True)
            error_msg = str(e)
            if "No /Root object" in error_msg or "really a PDF" in error_msg:
                raise RuntimeError(f"PDF conversion error: File '{pdf_path.name}' is not a valid PDF file. Please ensure the file is a properly formatted PDF document.")
            raise RuntimeError(f"PDF conversion error: {error_msg}")
        
        return ''.join(markdown_parts)

    def list_documents(self, session_id: str) -> List[Document]:
        """List all documents in a session."""
        doc_ids = self.sessions.get(session_id, [])
        return [self.documents[did] for did in doc_ids if did in self.documents]

    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a document by ID."""
        return self.documents.get(doc_id)

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document (allowed for ERROR, QUEUED, or CONVERTED status)."""
        doc = self.documents.get(doc_id)
        if not doc:
            return False
        # Allow deletion of ERROR, QUEUED, or CONVERTED documents
        allowed_statuses = ["ERROR", "QUEUED", "CONVERTED"]
        if doc.status not in allowed_statuses:
            raise ValueError(f"Can only delete documents with status: {', '.join(allowed_statuses)}")
        # Remove from session
        if doc.session_id in self.sessions:
            self.sessions[doc.session_id] = [d for d in self.sessions[doc.session_id] if d != doc_id]
        # Remove document
        del self.documents[doc_id]
        return True

    # -------- Chain Management --------
    def list_chains(self) -> List[Chain]:
        """List all available chains."""
        return list(self.chains.values())

    def get_chain(self, chain_version_id: str) -> Optional[Chain]:
        """Get a chain by version ID."""
        return self.chains.get(chain_version_id)

    def get_chain_by_id(self, chain_id: str) -> Optional[Chain]:
        """Get a chain by chain_id (not version_id)."""
        for chain in self.chains.values():
            if chain.chain_id == chain_id:
                return chain
        return None

    def create_chain(self, user_id: str, name: str, description: str, steps: List[Dict]) -> Chain:
        """Create a new chain."""
        # Validate chain structure
        is_valid, error_msg = self._validate_chain_structure(steps)
        if not is_valid:
            raise ValueError(f"Invalid chain structure: {error_msg}")

        # Ensure each step has model_config with defaults
        for step in steps:
            if "model_config" not in step:
                step["model_config"] = {
                    "model": "claude-3-haiku-20240307",
                    "max_tokens": 4096,
                    "temperature": 0.2
                }

        chain_id = f"chain_{uuid.uuid4().hex[:12]}"
        chain_version_id = self._generate_chain_version_id(chain_id)

        now = datetime.utcnow().isoformat() + "Z"

        chain = Chain(
            chain_id=chain_id,
            name=name,
            description=description,
            updated_at=now,
            chain_version_id=chain_version_id,
            step_count=len(steps),
            valid=is_valid,
            steps=steps,
        )

        self.chains[chain_version_id] = chain
        return chain

    def update_chain(self, chain_id: str, name: str, description: str, steps: List[Dict]) -> Chain:
        """Update an existing chain (mutates in place)."""
        # Find chain by chain_id
        chain = self.get_chain_by_id(chain_id)
        if not chain:
            raise ValueError(f"Chain not found: {chain_id}")

        # Validate new chain structure
        is_valid, error_msg = self._validate_chain_structure(steps)
        if not is_valid:
            raise ValueError(f"Invalid chain structure: {error_msg}")

        # Ensure each step has model_config with defaults
        for step in steps:
            if "model_config" not in step:
                step["model_config"] = {
                    "model": "claude-3-haiku-20240307",
                    "max_tokens": 4096,
                    "temperature": 0.2
                }

        # Update chain metadata and steps
        chain.name = name
        chain.description = description
        chain.step_count = len(steps)
        chain.steps = steps
        chain.valid = is_valid
        chain.updated_at = datetime.utcnow().isoformat() + "Z"

        return chain

    def _validate_chain_structure(self, steps: List[Dict]) -> tuple[bool, Optional[str]]:
        """
        Validate chain structure:
        - At least 1 step required
        - Step indices must be sequential (1, 2, 3...)
        - Each step must have required_inputs (array) and prompt (non-empty)
        - Inputs must be valid: R0 always valid, R1 only if step 1 exists, etc.
        """
        if not steps or len(steps) == 0:
            return False, "At least one step is required"

        # Check step indices are sequential starting at 1
        indices = [s.get("index") for s in steps]
        expected_indices = list(range(1, len(steps) + 1))
        if indices != expected_indices:
            return False, f"Step indices must be sequential starting at 1, got {indices}"

        # Validate each step
        for step in steps:
            index = step.get("index")
            required_inputs = step.get("required_inputs", [])
            prompt = step.get("prompt", "")

            # Check required_inputs is an array
            if not isinstance(required_inputs, list):
                return False, f"Step {index}: required_inputs must be an array"

            # Check at least one input
            if len(required_inputs) == 0:
                return False, f"Step {index}: At least one required_input is required"

            # Check prompt is non-empty
            if not prompt or not prompt.strip():
                return False, f"Step {index}: Prompt is required and cannot be empty"

            # Validate inputs: R0 always valid, R1 only if step 1 exists, etc.
            available_inputs = ["R0"]  # R0 always available
            for i in range(1, index):  # R1..R(N-1) available
                available_inputs.append(f"R{i}")

            invalid_inputs = [inp for inp in required_inputs if inp not in available_inputs]
            if invalid_inputs:
                return False, f"Step {index}: Invalid inputs {invalid_inputs}. Available: {available_inputs}"

        return True, None

    def _generate_chain_version_id(self, chain_id: str) -> str:
        """Generate a chain version ID."""
        # For now, simple format. Can enhance with versioning later.
        return f"cv_{chain_id}-v1"

    # -------- Run Management --------
    def create_run(self, session_id: str, chain_version_id: str) -> Run:
        """Create a new run for a session and chain version."""
        # Validate chain exists and is valid
        chain = self.chains.get(chain_version_id)
        if not chain:
            raise ValueError(f"Chain version not found: {chain_version_id}")
        if not chain.valid:
            raise ValueError(f"Chain is invalid: {chain_version_id}")

        # Validate all documents are CONVERTED
        docs = self.list_documents(session_id)
        if not docs:
            raise ValueError("No documents in session")
        unconverted = [d for d in docs if d.status != "CONVERTED"]
        if unconverted:
            raise ValueError(f"Some documents not converted: {[d.doc_id for d in unconverted]}")

        run_id = str(uuid.uuid4())
        run = Run(
            run_id=run_id,
            session_id=session_id,
            chain_version_id=chain_version_id,
            status="QUEUED",
            created_at=datetime.utcnow().isoformat() + "Z",
            document_ids=[d.doc_id for d in docs],
        )
        self.runs[run_id] = run

        # Trigger execution (for now: synchronous placeholder; later: async worker)
        self._trigger_run_execution(run_id)

        return run

    def _trigger_run_execution(self, run_id: str):
        """Trigger run execution with real Claude API calls."""
        run = self.runs.get(run_id)
        if not run:
            return

        run.status = "RUNNING"
        chain = self.chains.get(run.chain_version_id)
        if not chain:
            run.status = "ERROR"
            return

        if not LLM_AVAILABLE:
            logger.error("LLM client not available. Cannot execute steps.")
            run.status = "ERROR"
            return

        try:
            llm_client = get_llm_client()
            if not llm_client.is_available():
                raise RuntimeError("LLM client not initialized. Check ANTHROPIC_API_KEY.")

            # For each document, execute all steps sequentially
            for doc_id in run.document_ids:
                doc = self.documents.get(doc_id)
                if not doc or doc.status != "CONVERTED":
                    continue

                # Load R0 (converted document)
                if not doc.converted_md_path:
                    logger.warning(f"Document {doc_id} has no converted markdown")
                    continue

                r0_path = Path(doc.converted_md_path)
                if not r0_path.exists():
                    logger.warning(f"R0 file not found: {r0_path}")
                    continue

                r_outputs = {"R0": r0_path.read_text(encoding='utf-8')}

                # Execute each step in sequence
                for step_idx in range(1, chain.step_count + 1):
                    step_key = f"{run_id}:{doc_id}:{step_idx}"
                    
                    # Find step definition
                    step_def = next((s for s in chain.steps if s.get("index") == step_idx), None)
                    if not step_def:
                        continue

                    step_result = StepResult(
                        id=str(uuid.uuid4()),
                        run_id=run_id,
                        doc_id=doc_id,
                        step_index=step_idx,
                        status="RUNNING",
                    )
                    self.step_results[step_key] = step_result

                    try:
                        # Execute step with Claude API
                        output, tokens = self._execute_step(
                            llm_client=llm_client,
                            step_def=step_def,
                            r_outputs=r_outputs,
                            doc_id=doc_id,
                        )

                        # Store R(n) output
                        r_key = f"R{step_idx}"
                        r_outputs[r_key] = output

                        # Save R(n) artifact
                        run_dir = self.storage_base / "runs" / run_id / "docs" / doc_id
                        run_dir.mkdir(parents=True, exist_ok=True)
                        r_path = run_dir / f"{r_key}.md"
                        r_path.write_text(output, encoding='utf-8')

                        step_result.output_object_key = str(r_path)
                        step_result.status = "SUCCESS"
                        step_result.input_tokens = tokens.get("input_tokens", 0)
                        step_result.output_tokens = tokens.get("output_tokens", 0)
                        step_result.model = tokens.get("model", "claude-3-haiku-20240307")
                        
                        run.total_input_tokens += step_result.input_tokens
                        run.total_output_tokens += step_result.output_tokens

                    except Exception as e:
                        logger.error(f"Step {step_idx} execution failed for doc {doc_id}: {e}", exc_info=True)
                        step_result.status = "ERROR"
                        step_result.error_code = "EXECUTION_FAILED"
                        step_result.error_message = str(e)

        except Exception as e:
            logger.error(f"Run execution failed: {e}", exc_info=True)
            run.status = "ERROR"
            return

        # Check if all steps completed successfully
        all_complete = True
        for doc_id in run.document_ids:
            for step_idx in range(1, chain.step_count + 1):
                step_key = f"{run_id}:{doc_id}:{step_idx}"
                step_result = self.step_results.get(step_key)
                if step_result and step_result.status not in ("SUCCESS", "ERROR"):
                    all_complete = False
                    break
            if not all_complete:
                break

        run.status = "COMPLETE" if all_complete else "ERROR"

    def _execute_step(
        self,
        llm_client,
        step_def: Dict,
        r_outputs: Dict[str, str],
        doc_id: str,
    ) -> tuple[str, Dict]:
        """
        Execute a single step using Claude API.
        
        Returns:
            tuple: (output_text, tokens_dict)
        """
        step_prompt = step_def.get("prompt", "")
        required_inputs = step_def.get("required_inputs", [])
        
        # Build user message with labeled inputs
        user_message_parts = []
        user_message_parts.append(step_prompt)
        user_message_parts.append("\n\n---\n\n")
        user_message_parts.append("Available inputs:\n\n")
        
        for r_key in required_inputs:
            if r_key in r_outputs:
                user_message_parts.append(f"### {r_key}\n\n")
                user_message_parts.append(r_outputs[r_key])
                user_message_parts.append("\n\n---\n\n")
            else:
                raise ValueError(f"Required input {r_key} not found in available outputs")

        user_message = "".join(user_message_parts)

        # System prompt for step execution
        system_prompt = """You are executing a step in a multi-step document analysis chain.
Follow the step instructions precisely and produce the requested output.
The output should be well-formatted and ready for use in subsequent steps."""

        # Call Claude API with detailed response to get token usage
        start_time = time.time()
        
        # Use Anthropic API directly to get full response with usage info
        messages = [{"role": "user", "content": user_message}]
        response = llm_client.client.messages.create(
            model=llm_client.model,
            messages=messages,
            system=system_prompt,
            temperature=0.2,
            max_tokens=4096,
        )
        
        latency_ms = int((time.time() - start_time) * 1000)

        # Extract text from response
        response_text = ""
        for block in getattr(response, "content", []) or []:
            block_type = getattr(block, "type", None)
            if block_type == "text":
                block_text = getattr(block, "text", None)
                if isinstance(block_text, str):
                    response_text += block_text

        # Get token usage from response (Anthropic API includes this in response.usage)
        usage = getattr(response, "usage", None)
        input_tokens = 0
        output_tokens = 0
        
        if usage:
            input_tokens = getattr(usage, "input_tokens", 0) or 0
            output_tokens = getattr(usage, "output_tokens", 0) or 0

        tokens = {
            "input_tokens": input_tokens,
            "output_tokens": output_tokens,
            "model": llm_client.model or "claude-3-haiku-20240307",
            "latency_ms": latency_ms,
        }

        return response_text, tokens

    def get_run_progress(self, run_id: str) -> Dict:
        """Get run progress with per-document, per-step status."""
        run = self.runs.get(run_id)
        if not run:
            return None

        # Build progress rows per document
        rows = []
        for doc_id in run.document_ids:
            doc = self.documents.get(doc_id)
            if not doc:
                continue

            # Find highest completed step and aggregate status/tokens
            max_step = 0
            latest_status = "QUEUED"
            total_input_tokens = 0
            total_output_tokens = 0
            step_statuses = []

            chain = self.chains.get(run.chain_version_id)
            chain_step_count = chain.step_count if chain else 1

            for step_idx in range(1, chain_step_count + 1):
                step_key = f"{run_id}:{doc_id}:{step_idx}"
                step_result = self.step_results.get(step_key)
                if not step_result:
                    break
                max_step = step_idx
                step_statuses.append(step_result.status)
                latest_status = step_result.status
                if step_result.input_tokens:
                    total_input_tokens += step_result.input_tokens
                if step_result.output_tokens:
                    total_output_tokens += step_result.output_tokens

            step_label = f"R{max_step}" if max_step > 0 else "R0"

            # Determine doc status: SUCCESS if all steps are SUCCESS, ERROR if any ERROR, else RUNNING/QUEUED
            if all(s == "SUCCESS" for s in step_statuses) and len(step_statuses) == chain_step_count:
                doc_status = "SUCCESS"
                can_download = True
            elif any(s == "ERROR" for s in step_statuses):
                doc_status = "ERROR"
                can_download = False
            elif any(s == "RUNNING" for s in step_statuses):
                doc_status = "RUNNING"
                can_download = False
            else:
                doc_status = latest_status
                can_download = False

            rows.append({
                "doc_id": doc_id,
                "filename": doc.original_filename,
                "step_label": step_label,
                "status": doc_status,
                "input_tokens": total_input_tokens if total_input_tokens > 0 else None,
                "output_tokens": total_output_tokens if total_output_tokens > 0 else None,
                "can_download": can_download,
            })

        return {
            "run_id": run_id,
            "status": run.status,
            "rows": rows,
        }

    def get_download_url(self, run_id: str, doc_id: str) -> Optional[str]:
        """Get download URL for a document's final output."""
        return f"/api/bulk-doc-analysis/runs/{run_id}/download/{doc_id}"

    def get_final_output_path(self, run_id: str, doc_id: str, chain: Chain) -> Optional[Path]:
        """Get the path to the final step output (R(N)) for a document."""
        if chain.step_count == 0:
            return None
        
        final_r_key = f"R{chain.step_count}"
        run_dir = self.storage_base / "runs" / run_id / "docs" / doc_id
        r_path = run_dir / f"{final_r_key}.md"
        
        if r_path.exists():
            return r_path
        return None

