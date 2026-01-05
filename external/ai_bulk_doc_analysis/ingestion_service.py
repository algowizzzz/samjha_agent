"""
Ingestion Service - Handles file ingestion (programmatic and vision).
"""
import logging
import base64
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logging.warning("pdf2image not available")

try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    logging.warning("Pillow not available")

try:
    from external.platform.llm.client import get_llm_client
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    logging.warning("LLM client not available")

from .models import IngestionProfile
from .db_service import get_db_session

logger = logging.getLogger(__name__)

# Default vision prompt for document extraction
DEFAULT_VISION_PROMPT = """Extract all text from these document pages. Preserve the structure, formatting, and layout as much as possible. Include all tables, lists, and paragraphs. Convert the content to markdown format with appropriate headings, tables, and formatting."""


class IngestionService:
    """Service for file ingestion (programmatic and vision)."""
    
    def create_ingestion_profile(
        self,
        name: str,
        accepted_input_types: List[str],
        mode: str,
        vision_prompt: Optional[str] = None,
        vision_config: Optional[Dict] = None
    ) -> IngestionProfile:
        """
        Create an ingestion profile.
        
        Args:
            name: Profile name
            accepted_input_types: List of accepted file types ['PDF', 'DOCX', 'TXT', 'MD', 'CSV']
            mode: 'programmatic' or 'vision'
            vision_prompt: Vision prompt text (required if mode='vision')
            vision_config: Vision configuration dict with dpi, image_format, jpeg_quality, model, max_tokens, temperature
            
        Returns:
            Created IngestionProfile object
            
        Raises:
            ValueError: If validation fails
        """
        # Validate mode
        if mode not in ['programmatic', 'vision']:
            raise ValueError("Mode must be 'programmatic' or 'vision'")
        
        # Validate vision prompt
        if mode == 'vision' and not vision_prompt:
            raise ValueError("vision_prompt is required when mode='vision'")
        
        # Validate accepted types
        valid_types = ['PDF', 'DOCX', 'TXT', 'MD', 'CSV']
        for file_type in accepted_input_types:
            if file_type not in valid_types:
                raise ValueError(f"Invalid file type: {file_type}. Valid types: {valid_types}")
        
        # Set default vision_config if mode is vision and config not provided
        if mode == 'vision' and vision_config is None:
            vision_config = {
                "dpi": 200,
                "image_format": "PNG",
                "jpeg_quality": 85,
                "model": "claude-3-opus-20240229",
                "max_tokens": 4096,
                "temperature": 0.0
            }
        
        # Prepare metadata_json
        metadata = {}
        if vision_config:
            metadata["vision_config"] = vision_config
        
        with get_db_session() as db:
            ingestion_profile_id = f"ing_{uuid.uuid4().hex[:12]}"
            profile = IngestionProfile(
                ingestion_profile_id=ingestion_profile_id,
                name=name,
                accepted_input_types=accepted_input_types,
                mode=mode,
                vision_prompt=vision_prompt,
                metadata_json=metadata if metadata else {}
            )
            db.add(profile)
            db.commit()
            
            logger.info(f"Created ingestion profile {ingestion_profile_id}")
            # Return a simple object with the data (avoiding session binding issues)
            class ProfileResult:
                pass
            result = ProfileResult()
            result.ingestion_profile_id = ingestion_profile_id
            result.name = name
            result.accepted_input_types = accepted_input_types
            result.mode = mode
            result.vision_prompt = vision_prompt
            return result
    
    def get_ingestion_profile(self, ingestion_profile_id: str) -> Optional[IngestionProfile]:
        """Get ingestion profile by ID."""
        with get_db_session() as db:
            return db.query(IngestionProfile).filter(
                IngestionProfile.ingestion_profile_id == ingestion_profile_id
            ).first()
    
    def list_ingestion_profiles(self) -> List[IngestionProfile]:
        """List all ingestion profiles."""
        with get_db_session() as db:
            return db.query(IngestionProfile).all()
    
    def ingest_file(
        self,
        file_path: Path,
        ingestion_profile: IngestionProfile
    ) -> Tuple[str, Dict]:
        """
        Ingest a file and return R0 content.
        
        Args:
            file_path: Path to file
            ingestion_profile: IngestionProfile object
            
        Returns:
            Tuple of (r0_content, metadata_dict)
            
        Raises:
            ValueError: If file type not supported or mode invalid
            RuntimeError: If ingestion fails
        """
        file_type = self._detect_file_type(file_path)
        
        # Validate file type
        if file_type not in ingestion_profile.accepted_input_types:
            raise ValueError(f"File type {file_type} not in accepted types: {ingestion_profile.accepted_input_types}")
        
        # Route to appropriate ingestion method
        if ingestion_profile.mode == 'programmatic':
            return self._ingest_programmatic(file_path, file_type)
        elif ingestion_profile.mode == 'vision':
            if file_type != 'PDF':
                raise ValueError("Vision ingestion only supports PDF files")
            return self._ingest_vision(file_path, ingestion_profile)
        else:
            raise ValueError(f"Unknown ingestion mode: {ingestion_profile.mode}")
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type from extension."""
        ext = file_path.suffix.lower()
        mapping = {
            '.pdf': 'PDF',
            '.docx': 'DOCX',
            '.txt': 'TXT',
            '.md': 'MD',
            '.csv': 'CSV'
        }
        return mapping.get(ext, 'UNKNOWN')
    
    def _ingest_programmatic(self, file_path: Path, file_type: str) -> Tuple[str, Dict]:
        """
        Programmatic ingestion - extract text using libraries.
        
        Returns:
            Tuple of (r0_content, metadata)
        """
        metadata = {"method": "programmatic", "file_type": file_type}
        
        if file_type == 'PDF':
            if not PDFPLUMBER_AVAILABLE:
                raise RuntimeError("pdfplumber not installed. Install with: pip install pdfplumber")
            
            markdown_parts = []
            markdown_parts.append(f"# Document: {file_path.name}\n\n")
            
            # Validate file is actually a PDF by checking magic bytes
            try:
                with open(file_path, 'rb') as f:
                    header = f.read(4)
                    if header != b'%PDF':
                        raise RuntimeError(f"File '{file_path.name}' does not appear to be a valid PDF file. PDF files must start with '%PDF' magic bytes. This file appears to be a different format.")
            except Exception as e:
                if "magic bytes" in str(e) or "does not appear" in str(e):
                    raise
                # If it's a different error, let pdfplumber handle it
                pass
            
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        markdown_parts.append(f"## Page {page_num}\n\n")
                        
                        # Extract text
                        text = page.extract_text()
                        if text:
                            text = text.strip()
                            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                            markdown_parts.append('\n\n'.join(paragraphs))
                            markdown_parts.append('\n\n')
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table_num, table in enumerate(tables, start=1):
                                markdown_parts.append(f"### Table {table_num} (Page {page_num})\n\n")
                                if table and len(table) > 0:
                                    if table[0]:
                                        header = "| " + " | ".join(str(cell or "") for cell in table[0]) + " |"
                                        separator = "| " + " | ".join("---" for _ in table[0]) + " |"
                                        markdown_parts.append(header)
                                        markdown_parts.append(separator)
                                    
                                    for row in table[1:]:
                                        if row:
                                            row_md = "| " + " | ".join(str(cell or "") for cell in row) + " |"
                                            markdown_parts.append(row_md)
                                    markdown_parts.append('\n\n')
                
                r0_content = ''.join(markdown_parts)
                metadata["page_count"] = len(pdf.pages) if 'pdf' in locals() else 0
                return r0_content, metadata
                
            except RuntimeError:
                # Re-raise validation errors as-is
                raise
            except Exception as e:
                logger.error(f"Error converting PDF {file_path}: {e}", exc_info=True)
                # Provide more helpful error message
                error_msg = str(e)
                if "No /Root object" in error_msg or "really a PDF" in error_msg:
                    raise RuntimeError(f"PDF conversion error: File '{file_path.name}' is not a valid PDF file. Please ensure the file is a properly formatted PDF document.")
                raise RuntimeError(f"PDF conversion error: {error_msg}")
        
        elif file_type == 'DOCX':
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
            
            markdown_parts = []
            markdown_parts.append(f"# Document: {file_path.name}\n\n")
            
            try:
                doc = DocxDocument(file_path)
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        # Simple heading detection (if style indicates heading)
                        if para.style.name.startswith('Heading'):
                            level = para.style.name.replace('Heading ', '')
                            markdown_parts.append(f"{'#' * int(level)} {para.text}\n\n")
                        else:
                            markdown_parts.append(f"{para.text}\n\n")
                
                # Extract tables
                for table_num, table in enumerate(doc.tables, start=1):
                    markdown_parts.append(f"### Table {table_num}\n\n")
                    if table.rows:
                        # Header row
                        header_row = table.rows[0]
                        header = "| " + " | ".join(str(cell.text or "") for cell in header_row.cells) + " |"
                        separator = "| " + " | ".join("---" for _ in header_row.cells) + " |"
                        markdown_parts.append(header)
                        markdown_parts.append(separator)
                        
                        # Data rows
                        for row in table.rows[1:]:
                            row_md = "| " + " | ".join(str(cell.text or "") for cell in row.cells) + " |"
                            markdown_parts.append(row_md)
                        markdown_parts.append('\n\n')
                
                r0_content = ''.join(markdown_parts)
                metadata["paragraph_count"] = len(doc.paragraphs)
                return r0_content, metadata
                
            except Exception as e:
                logger.error(f"Error converting DOCX {file_path}: {e}", exc_info=True)
                raise RuntimeError(f"DOCX conversion error: {str(e)}")
        
        elif file_type in ['TXT', 'MD']:
            # Direct read
            try:
                r0_content = file_path.read_text(encoding='utf-8')
                metadata["char_count"] = len(r0_content)
                return r0_content, metadata
            except Exception as e:
                logger.error(f"Error reading {file_type} {file_path}: {e}", exc_info=True)
                raise RuntimeError(f"File read error: {str(e)}")
        
        elif file_type == 'CSV':
            # CSV is handled specially - returns metadata only, actual ingestion creates tasks
            # This method is called but CSV should be handled in conversion worker
            try:
                import pandas as pd
                df = pd.read_csv(file_path)
                metadata["row_count"] = len(df)
                metadata["column_count"] = len(df.columns)
                metadata["columns"] = df.columns.tolist()
                # Return empty content - CSV ingestion creates tasks, not R0
                return "", metadata
            except Exception as e:
                logger.error(f"Error reading CSV {file_path}: {e}", exc_info=True)
                raise RuntimeError(f"CSV read error: {str(e)}")
        
        else:
            raise ValueError(f"Unsupported file type for programmatic ingestion: {file_type}")
    
    def _ingest_vision(self, file_path: Path, ingestion_profile: IngestionProfile) -> Tuple[str, Dict]:
        """
        Vision ingestion - convert PDF to images and call Claude Vision API per page.
        
        Args:
            file_path: Path to PDF file
            ingestion_profile: IngestionProfile object with vision config
            
        Returns:
            Tuple of (r0_content, metadata)
        """
        if not PDF2IMAGE_AVAILABLE:
            raise RuntimeError("pdf2image not installed. Install with: pip install pdf2image")
        
        if not PILLOW_AVAILABLE:
            raise RuntimeError("Pillow not installed. Install with: pip install Pillow")
        
        if not LLM_AVAILABLE:
            raise RuntimeError("LLM client not available")
        
        # Extract vision configuration from metadata_json
        vision_config = {}
        if ingestion_profile.metadata_json and isinstance(ingestion_profile.metadata_json, dict):
            vision_config = ingestion_profile.metadata_json.get("vision_config", {})
        
        # Get config values with defaults
        dpi = vision_config.get("dpi", 200)
        image_format = vision_config.get("image_format", "PNG").upper()
        jpeg_quality = vision_config.get("jpeg_quality", 85)
        model = vision_config.get("model", "claude-3-opus-20240229")
        max_tokens = vision_config.get("max_tokens", 4096)
        temperature = vision_config.get("temperature", 0.0)
        vision_prompt = ingestion_profile.vision_prompt or DEFAULT_VISION_PROMPT
        
        metadata = {
            "method": "vision",
            "file_type": "PDF",
            "dpi": dpi,
            "image_format": image_format,
            "model": model
        }
        
        try:
            import io
            
            # Convert PDF pages to images with configurable DPI
            images = convert_from_path(str(file_path), dpi=dpi, fmt=image_format.lower())
            metadata["page_count"] = len(images)
            
            # Initialize LLM client
            llm_client = get_llm_client()
            if not llm_client.is_available():
                raise RuntimeError("LLM client not initialized. Check ANTHROPIC_API_KEY.")
            
            # Process one page at a time
            r0_parts = []
            r0_parts.append(f"# Document: {file_path.name}\n\n")
            total_input_tokens = 0
            total_output_tokens = 0
            failed_pages = []
            
            for page_num, img in enumerate(images, start=1):
                try:
                    # Encode single image with configurable format
                    img_buffer = io.BytesIO()
                    if image_format == "JPEG":
                        img.save(img_buffer, format='JPEG', quality=jpeg_quality, optimize=True)
                        media_type = "image/jpeg"
                    else:
                        img.save(img_buffer, format='PNG', optimize=True)
                        media_type = "image/png"
                    
                    img_bytes = img_buffer.getvalue()
                    img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                    
                    # Create messages for single page
                    messages = [{
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": img_base64
                        }
                    }, {
                        "type": "text",
                        "text": vision_prompt
                    }]
                    
                    # Make API call for this page
                    response = llm_client.client.messages.create(
                        model=model,
                        messages=[{"role": "user", "content": messages}],
                        max_tokens=max_tokens,
                        temperature=temperature
                    )
                    
                    # Extract text from response
                    page_content = ""
                    for block in getattr(response, "content", []) or []:
                        block_type = getattr(block, "type", None)
                        if block_type == "text":
                            block_text = getattr(block, "text", None)
                            if isinstance(block_text, str):
                                page_content += block_text
                    
                    # Add page marker and content
                    r0_parts.append(f"## Page {page_num}\n\n")
                    r0_parts.append(page_content)
                    r0_parts.append("\n\n")
                    
                    # Track tokens
                    usage = getattr(response, "usage", None)
                    if usage:
                        page_input_tokens = getattr(usage, "input_tokens", 0) or 0
                        page_output_tokens = getattr(usage, "output_tokens", 0) or 0
                        total_input_tokens += page_input_tokens
                        total_output_tokens += page_output_tokens
                    
                    logger.info(f"Processed page {page_num}/{len(images)} of {file_path.name}")
                    
                except Exception as e:
                    logger.error(f"Error processing page {page_num} of {file_path.name}: {e}", exc_info=True)
                    failed_pages.append(page_num)
                    # Add error marker but continue processing
                    r0_parts.append(f"## Page {page_num}\n\n")
                    r0_parts.append(f"*Error processing this page: {str(e)}*\n\n")
            
            # Combine all pages
            r0_content = ''.join(r0_parts)
            
            # Add metadata
            metadata["input_tokens"] = total_input_tokens
            metadata["output_tokens"] = total_output_tokens
            if failed_pages:
                metadata["failed_pages"] = failed_pages
                metadata["warning"] = f"Failed to process {len(failed_pages)} page(s): {failed_pages}"
            
            logger.info(f"Vision ingestion completed for {file_path.name}, {metadata.get('page_count', 0)} pages, {total_input_tokens} input tokens, {total_output_tokens} output tokens")
            return r0_content, metadata
            
        except Exception as e:
            logger.error(f"Vision ingestion failed for {file_path}: {e}", exc_info=True)
            raise RuntimeError(f"Vision ingestion error: {str(e)}")
    
    def estimate_ingestion_tokens(
        self,
        file_path: Path,
        ingestion_profile: IngestionProfile
    ) -> int:
        """
        Estimate token count for ingestion.
        
        Args:
            file_path: Path to file
            ingestion_profile: IngestionProfile object
            
        Returns:
            Estimated token count
        """
        if ingestion_profile.mode == 'programmatic':
            # For programmatic: estimate based on file size
            # Rough estimate: 1 token ≈ 4 characters
            file_size = file_path.stat().st_size
            estimated_chars = file_size  # Conservative estimate
            return estimated_chars // 4
        
        elif ingestion_profile.mode == 'vision':
            # For vision: prompt tokens + image payload estimate per page
            # Higher DPI = more tokens per image
            try:
                # Extract vision_config to get DPI
                vision_config = {}
                if ingestion_profile.metadata_json and isinstance(ingestion_profile.metadata_json, dict):
                    vision_config = ingestion_profile.metadata_json.get("vision_config", {})
                dpi = vision_config.get("dpi", 200)
                
                # Higher DPI = more tokens per image
                # Rough estimate: 200 DPI ≈ 1000 tokens, 300 DPI ≈ 1500 tokens
                tokens_per_image = int(1000 * (dpi / 200))
                
                if PDF2IMAGE_AVAILABLE:
                    # Get actual page count
                    images = convert_from_path(str(file_path), first_page=1, last_page=1, dpi=dpi)
                    sample_image_count = len(images) if images else 1
                    # Estimate total pages (rough: assume similar to sample)
                    # For better accuracy, could count all pages, but this is just estimation
                    total_pages = sample_image_count  # Conservative: use sample count
                else:
                    # Fallback: estimate 1 page
                    total_pages = 1
                
                # Estimate prompt tokens (rough: 1 token ≈ 4 chars)
                prompt_tokens = len(ingestion_profile.vision_prompt or "") // 4
                # Estimate image tokens per page, multiplied by page count (since we process per-page)
                image_tokens_per_page = tokens_per_image
                total_image_tokens = image_tokens_per_page * total_pages
                
                # Total: prompt tokens (repeated per page) + image tokens per page * page count
                # Actually, prompt is sent once per page, so multiply prompt tokens by page count too
                return (prompt_tokens * total_pages) + total_image_tokens
            except Exception:
                # Fallback estimate
                return 5000

