"""
Output Conversion Worker - Converts markdown output to PDF/DOCX/CSV format.
This runs after all AI steps complete for a document.
"""
import logging
from pathlib import Path
from typing import Dict, Any, Optional

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab not available")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas not available")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logging.warning("openpyxl not available")

logger = logging.getLogger(__name__)


def convert_output_format_job(job_data: Dict[str, Any], storage_base: Path):
    """
    RQ job handler for converting markdown output to PDF/DOCX/CSV format.
    
    Args:
        job_data: {
            "run_id": str,
            "doc_id": str,
            "export_format": str,  # 'PDF', 'DOCX', 'CSV'
            "markdown_path": str,  # Relative path to the markdown file
            "idempotency_key": str
        }
        storage_base: Base path for storage
    """
    run_id = job_data["run_id"]
    doc_id = job_data["doc_id"]
    export_format = job_data["export_format"]
    markdown_path = Path(job_data["markdown_path"])
    
    logger.info(f"Converting output for run {run_id}, doc {doc_id} to {export_format}")
    
    try:
        # Resolve the markdown file path
        if markdown_path.is_absolute():
            md_full_path = markdown_path
        else:
            md_full_path = storage_base / markdown_path
        
        if not md_full_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_full_path}")
        
        # Read markdown content
        md_content = md_full_path.read_text(encoding='utf-8')
        
        # Determine output path (same directory as markdown, different extension)
        output_dir = md_full_path.parent
        base_name = md_full_path.stem  # Remove .md extension
        
        if export_format == 'PDF':
            if not REPORTLAB_AVAILABLE:
                raise RuntimeError("reportlab not installed. Install with: pip install reportlab")
            
            output_path = output_dir / f"{base_name}.pdf"
            _convert_md_to_pdf(md_content, output_path)
            
        elif export_format == 'DOCX':
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
            
            output_path = output_dir / f"{base_name}.docx"
            _convert_md_to_docx(md_content, output_path)
            
        elif export_format == 'CSV':
            if not PANDAS_AVAILABLE:
                raise RuntimeError("pandas not installed. Install with: pip install pandas")
            
            output_path = output_dir / f"{base_name}.csv"
            _convert_md_to_csv(md_content, output_path, base_name, output_dir)
            
        elif export_format == 'XLSX':
            if not OPENPYXL_AVAILABLE:
                raise RuntimeError("openpyxl not installed. Install with: pip install openpyxl")
            if not PANDAS_AVAILABLE:
                raise RuntimeError("pandas not installed. Install with: pip install pandas")
            
            output_path = output_dir / f"{base_name}.xlsx"
            _convert_md_to_xlsx(md_content, output_path)
            
        else:
            raise ValueError(f"Unsupported export format: {export_format}")
        
        logger.info(f"Successfully converted output to {export_format}: {output_path}")
        return {
            "status": "SUCCESS",
            "output_path": str(output_path.relative_to(storage_base)),
        }
        
    except Exception as e:
        logger.error(f"Output conversion failed: {e}", exc_info=True)
        raise


def _convert_md_to_pdf(md_content: str, output_path: Path):
    """Convert markdown content to PDF."""
    doc = SimpleDocTemplate(str(output_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Simple markdown to PDF conversion
    # Split by lines and process
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        
        # Simple heading detection
        if line.startswith('# '):
            story.append(Paragraph(line[2:], styles['Heading1']))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading3']))
        else:
            # Regular paragraph
            # Escape HTML/XML special chars for reportlab
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(escaped, styles['Normal']))
    
    doc.build(story)


def _convert_md_to_docx(md_content: str, output_path: Path):
    """Convert markdown content to DOCX."""
    doc = DocxDocument()
    
    # Simple markdown to DOCX conversion
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Simple heading detection
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    doc.save(str(output_path))


def _extract_tables_from_markdown(md_content: str) -> list:
    """Extract all markdown tables from content.
    
    Returns:
        List of tables, each table is a list of rows (each row is a list of cells)
    """
    lines = md_content.split('\n')
    tables = []
    current_table = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip empty lines - they separate tables
        if not line_stripped:
            if current_table:
                tables.append(current_table)
                current_table = []
            continue
        
        # Skip markdown table separator rows (|---|---|)
        if line_stripped.startswith('|---'):
            continue
        
        # Detect markdown table rows (containing |)
        if '|' in line_stripped:
            # Split by | and clean up
            cells = [cell.strip() for cell in line_stripped.split('|')]
            # Remove empty first/last cells from leading/trailing |
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:
                current_table.append(cells)
        else:
            # Non-table line - if we have a table, save it
            if current_table:
                tables.append(current_table)
                current_table = []
    
    # Don't forget the last table
    if current_table:
        tables.append(current_table)
    
    return tables


def _convert_md_to_csv(md_content: str, output_path: Path, base_name: str, output_dir: Path):
    """Convert markdown table content to CSV.
    
    If multiple tables are found:
    - Creates a ZIP file containing multiple CSV files (one per table)
    - Otherwise creates a single CSV file
    """
    import zipfile
    
    tables = _extract_tables_from_markdown(md_content)
    
    if not tables:
        # No table found, create single-column CSV with content
        df = pd.DataFrame({'content': md_content.split('\n')})
        df.to_csv(output_path, index=False)
        return
    
    if len(tables) == 1:
        # Single table: create one CSV
        rows = tables[0]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0] if rows else [])
            df.to_csv(output_path, index=False)
    else:
        # Multiple tables: create ZIP with multiple CSVs
        zip_path = output_dir / f"{base_name}.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, table in enumerate(tables, 1):
                if table:
                    # Create CSV for this table
                    table_df = pd.DataFrame(table[1:], columns=table[0] if table else [])
                    csv_filename = f"{base_name}_table_{i}.csv"
                    csv_path = output_dir / csv_filename
                    table_df.to_csv(csv_path, index=False)
                    # Add to ZIP
                    zip_file.write(csv_path, csv_filename)
                    # Remove temp CSV file
                    csv_path.unlink()
        
        # Remove single CSV file if it exists, keep ZIP
        output_path.unlink(missing_ok=True)


def _convert_md_to_xlsx(md_content: str, output_path: Path):
    """Convert markdown table content to Excel with multiple sheets (one per table)."""
    tables = _extract_tables_from_markdown(md_content)
    
    if not tables:
        # No table found, create single sheet with content
        df = pd.DataFrame({'content': md_content.split('\n')})
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Sheet1', index=False)
        return
    
    # Create Excel with multiple sheets (one per table)
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for i, table in enumerate(tables, 1):
            if table:
                # Create DataFrame for this table
                table_df = pd.DataFrame(table[1:], columns=table[0] if table else [])
                # Sheet names must be <= 31 chars and no special chars
                sheet_name = f"Table_{i}"[:31]
                table_df.to_excel(writer, sheet_name=sheet_name, index=False)

