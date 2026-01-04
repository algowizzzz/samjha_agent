"""
Export Service - Handles result export in multiple formats.
"""
import logging
import uuid
import json
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas not available")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab not available")

from .models import ExportProfile, Run, StepResult, ExecutionTask
from .db_service import get_db_session

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting workflow results."""
    
    def create_export_profile(
        self,
        name: str,
        format: str,
        config_json: Optional[Dict] = None
    ) -> ExportProfile:
        """
        Create an export profile.
        
        Args:
            name: Profile name
            format: Export format ('CSV', 'JSON', 'MD', 'DOCX', 'PDF')
            config_json: Optional configuration dict
            
        Returns:
            Created ExportProfile object
            
        Raises:
            ValueError: If validation fails
        """
        valid_formats = ['CSV', 'JSON', 'MD', 'DOCX', 'PDF', 'XLSX']
        if format not in valid_formats:
            raise ValueError(f"Format must be one of: {valid_formats}")
        
        with get_db_session() as db:
            export_profile_id = f"exp_{uuid.uuid4().hex[:12]}"
            profile = ExportProfile(
                export_profile_id=export_profile_id,
                name=name,
                format=format,
                config_json=config_json or {}
            )
            db.add(profile)
            db.commit()
            
            logger.info(f"Created export profile {export_profile_id}")
            # Return a simple object with the data (avoiding session binding issues)
            class ProfileResult:
                pass
            result = ProfileResult()
            result.export_profile_id = export_profile_id
            result.name = name
            result.format = format
            result.config_json = config_json or {}
            return result
    
    def get_export_profile(self, export_profile_id: str) -> Optional[ExportProfile]:
        """Get export profile by ID."""
        with get_db_session() as db:
            return db.query(ExportProfile).filter(
                ExportProfile.export_profile_id == export_profile_id
            ).first()
    
    def list_export_profiles(self) -> List[ExportProfile]:
        """List all export profiles."""
        with get_db_session() as db:
            return db.query(ExportProfile).all()
    
    def export_results(
        self,
        run_id: str,
        export_profile: ExportProfile,
        storage_base: Path
    ) -> Path:
        """
        Export run results in the specified format.
        
        Args:
            run_id: Run ID
            export_profile: ExportProfile object
            storage_base: Base path for storage
            
        Returns:
            Path to exported file
            
        Raises:
            ValueError: If format not supported
            RuntimeError: If export fails
        """
        format = export_profile.format
        
        if format == 'CSV':
            return self._export_csv(run_id, storage_base)
        elif format == 'JSON':
            return self._export_json(run_id, storage_base)
        elif format == 'MD':
            return self._export_markdown(run_id, storage_base)
        elif format == 'DOCX':
            return self._export_docx(run_id, storage_base)
        elif format == 'PDF':
            return self._export_pdf(run_id, storage_base)
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    def _export_csv(self, run_id: str, storage_base: Path) -> Path:
        """Export results as CSV."""
        if not PANDAS_AVAILABLE:
            raise RuntimeError("pandas not installed. Install with: pip install pandas")
        
        with get_db_session() as db:
            # Check if CSV workflow (has execution_tasks)
            tasks = db.query(ExecutionTask).filter(ExecutionTask.run_id == run_id).all()
            
            if tasks:
                # CSV workflow: compile from tasks
                rows = []
                for task in tasks:
                    # Get final step result for this task
                    final_step = (
                        db.query(StepResult)
                        .filter(
                            StepResult.run_id == run_id,
                            StepResult.task_id == task.task_id
                        )
                        .order_by(StepResult.step_index.desc())
                        .first()
                    )
                    
                    if final_step and final_step.status == 'SUCCESS':
                        # Load output
                        output_path = storage_base / final_step.output_object_key if final_step.output_object_key else None
                        if output_path and output_path.exists():
                            output_text = output_path.read_text(encoding='utf-8')
                        else:
                            output_text = ""
                    else:
                        output_text = ""
                    
                    # Combine original row data with output
                    row_data = dict(task.row_data) if task.row_data else {}
                    row_data['_output'] = output_text
                    row_data['_row_index'] = task.row_index
                    rows.append(row_data)
                
                # Create DataFrame
                df = pd.DataFrame(rows)
                
                # Export to CSV
                output_path = storage_base / "runs" / run_id / "export.csv"
                output_path.parent.mkdir(parents=True, exist_ok=True)
                df.to_csv(output_path, index=False)
                
                return output_path
            else:
                # Non-CSV workflow: compile from documents
                step_results = (
                    db.query(StepResult)
                    .filter(StepResult.run_id == run_id, StepResult.status == 'SUCCESS')
                    .order_by(StepResult.doc_id, StepResult.step_index)
                    .all()
                )
                
                rows = []
                current_doc = None
                current_row = {}
                
                for sr in step_results:
                    if sr.doc_id != current_doc:
                        if current_row:
                            rows.append(current_row)
                        current_doc = sr.doc_id
                        current_row = {'doc_id': sr.doc_id, 'step_index': sr.step_index}
                    
                    # Load output
                    output_path = storage_base / sr.output_object_key if sr.output_object_key else None
                    if output_path and output_path.exists():
                        output_text = output_path.read_text(encoding='utf-8')
                    else:
                        output_text = ""
                    
                    current_row[f'step_{sr.step_index}_output'] = output_text
                
                if current_row:
                    rows.append(current_row)
                
                df = pd.DataFrame(rows)
                output_path = storage_base / "runs" / run_id / "export.csv"
                output_path.parent.mkdir(parents=True, exist_ok=True)
                df.to_csv(output_path, index=False)
                
                return output_path
    
    def _export_json(self, run_id: str, storage_base: Path) -> Path:
        """Export results as JSON."""
        with get_db_session() as db:
            # Get all step results
            step_results = (
                db.query(StepResult)
                .filter(StepResult.run_id == run_id, StepResult.status == 'SUCCESS')
                .order_by(StepResult.doc_id, StepResult.step_index)
                .all()
            )
            
            results = []
            for sr in step_results:
                # Load output
                output_path = storage_base / sr.output_object_key if sr.output_object_key else None
                output_text = ""
                if output_path and output_path.exists():
                    output_text = output_path.read_text(encoding='utf-8')
                
                results.append({
                    "doc_id": sr.doc_id,
                    "task_id": sr.task_id,
                    "step_index": sr.step_index,
                    "output": output_text,
                    "input_tokens": sr.input_tokens,
                    "output_tokens": sr.output_tokens,
                })
            
            # Export to JSON
            output_path = storage_base / "runs" / run_id / "export.json"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2)
            
            return output_path
    
    def _export_markdown(self, run_id: str, storage_base: Path) -> Path:
        """Export results as Markdown (zip if multiple files)."""
        with get_db_session() as db:
            # Get all step results grouped by doc
            step_results = (
                db.query(StepResult)
                .filter(StepResult.run_id == run_id, StepResult.status == 'SUCCESS')
                .order_by(StepResult.doc_id, StepResult.step_index)
                .all()
            )
            
            files = []
            current_doc = None
            md_parts = []
            
            for sr in step_results:
                if sr.doc_id != current_doc:
                    if current_doc and md_parts:
                        # Save previous doc's markdown
                        md_path = storage_base / "runs" / run_id / f"export_{current_doc}.md"
                        md_path.parent.mkdir(parents=True, exist_ok=True)
                        md_path.write_text(''.join(md_parts), encoding='utf-8')
                        files.append(md_path)
                    
                    current_doc = sr.doc_id
                    md_parts = [f"# Document: {current_doc}\n\n"]
                
                # Load output
                output_path = storage_base / sr.output_object_key if sr.output_object_key else None
                if output_path and output_path.exists():
                    output_text = output_path.read_text(encoding='utf-8')
                    md_parts.append(f"## Step {sr.step_index}\n\n{output_text}\n\n")
            
            # Save last doc
            if current_doc and md_parts:
                md_path = storage_base / "runs" / run_id / f"export_{current_doc}.md"
                md_path.parent.mkdir(parents=True, exist_ok=True)
                md_path.write_text(''.join(md_parts), encoding='utf-8')
                files.append(md_path)
            
            # If multiple files, create zip
            if len(files) > 1:
                zip_path = storage_base / "runs" / run_id / "export.zip"
                with zipfile.ZipFile(zip_path, 'w') as zf:
                    for f in files:
                        zf.write(f, f.name)
                return zip_path
            elif len(files) == 1:
                return files[0]
            else:
                # No results, create empty file
                output_path = storage_base / "runs" / run_id / "export.md"
                output_path.parent.mkdir(parents=True, exist_ok=True)
                output_path.write_text("# No results\n\n", encoding='utf-8')
                return output_path
    
    def _export_docx(self, run_id: str, storage_base: Path) -> Path:
        """Export results as DOCX."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        doc = DocxDocument()
        doc.add_heading('Workflow Results', 0)
        
        with get_db_session() as db:
            step_results = (
                db.query(StepResult)
                .filter(StepResult.run_id == run_id, StepResult.status == 'SUCCESS')
                .order_by(StepResult.doc_id, StepResult.step_index)
                .all()
            )
            
            current_doc = None
            for sr in step_results:
                if sr.doc_id != current_doc:
                    if current_doc:
                        doc.add_page_break()
                    current_doc = sr.doc_id
                    doc.add_heading(f'Document: {current_doc}', 1)
                
                doc.add_heading(f'Step {sr.step_index}', 2)
                
                # Load output
                output_path = storage_base / sr.output_object_key if sr.output_object_key else None
                if output_path and output_path.exists():
                    output_text = output_path.read_text(encoding='utf-8')
                    # Simple paragraph addition (could be enhanced with markdown parsing)
                    doc.add_paragraph(output_text)
        
        output_path = storage_base / "runs" / run_id / "export.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
    
    def _export_pdf(self, run_id: str, storage_base: Path) -> Path:
        """Export results as PDF."""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("reportlab not installed. Install with: pip install reportlab")
        
        output_path = storage_base / "runs" / run_id / "export.pdf"
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        story.append(Paragraph("Workflow Results", styles['Title']))
        story.append(Spacer(1, 12))
        
        with get_db_session() as db:
            step_results = (
                db.query(StepResult)
                .filter(StepResult.run_id == run_id, StepResult.status == 'SUCCESS')
                .order_by(StepResult.doc_id, StepResult.step_index)
                .all()
            )
            
            current_doc = None
            for sr in step_results:
                if sr.doc_id != current_doc:
                    if current_doc:
                        story.append(Spacer(1, 12))
                    current_doc = sr.doc_id
                    story.append(Paragraph(f"Document: {current_doc}", styles['Heading1']))
                
                story.append(Paragraph(f"Step {sr.step_index}", styles['Heading2']))
                
                # Load output
                output_path_obj = storage_base / sr.output_object_key if sr.output_object_key else None
                if output_path_obj and output_path_obj.exists():
                    output_text = output_path_obj.read_text(encoding='utf-8')
                    # Simple text addition (could be enhanced)
                    story.append(Paragraph(output_text, styles['Normal']))
        
        doc.build(story)
        return output_path

